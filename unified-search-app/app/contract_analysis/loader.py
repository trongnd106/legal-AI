# Copyright (c) 2024 Microsoft Corporation.
# Licensed under the MIT License

"""Load PDF / DOCX / TXT và kiểm tra từ khóa HĐLĐ."""

from __future__ import annotations

import re
from pathlib import Path

from contract_analysis.constants import LABOR_CONTRACT_KEYWORDS
from contract_analysis.schema import ContractDocument, ContractMetadata


def score_labor_keywords(text: str) -> float:
    if not text.strip():
        return 0.0
    lower = text.lower()
    hits = sum(1 for kw in LABOR_CONTRACT_KEYWORDS if kw.lower() in lower)
    return min(1.0, hits / max(3, len(LABOR_CONTRACT_KEYWORDS) * 0.35))


def _load_pdf_text_layer(path: Path) -> tuple[str, list[dict[str, str | int]], int | None]:
    import pdfplumber

    parts: list[str] = []
    page_records: list[dict[str, str | int]] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            parts.append(t)
            page_records.append({"page_num": i + 1, "text": t})
        n_pages = len(pdf.pages)
    return "\n\n".join(parts), page_records, n_pages


def _pdf_looks_scanned(
    parts: list[str],
    n_pages: int,
    *,
    min_total_chars: int = 120,
    min_avg_per_page: float = 35.0,
) -> bool:
    """Heuristic: ít ký tự trích từ text layer → có thể là PDF scan."""
    if n_pages <= 0:
        return True
    total = sum(len((p or "").strip()) for p in parts)
    avg = total / n_pages
    return total < min_total_chars or avg < min_avg_per_page


def _load_pdf(
    path: Path,
    *,
    force_ocr: bool = False,
    detect_scan: bool = True,
) -> tuple[str, list[dict[str, str | int]], int | None, str]:
    """Trả ``(raw, pages, n_pages, extraction_tag)`` với tag ``pdfplumber`` hoặc ``paddleocr``."""
    raw_tb, page_tb, n_pages = _load_pdf_text_layer(path)
    use_ocr = force_ocr or (
        detect_scan and _pdf_looks_scanned([p["text"] for p in page_tb], int(n_pages or 0))
    )
    if use_ocr:
        from contract_analysis.ocr_pdf import extract_pdf_text_paddleocr

        raw_ocr, page_ocr, n_ocr = extract_pdf_text_paddleocr(path)
        return raw_ocr, page_ocr, n_ocr, "paddleocr"
    return raw_tb, page_tb, n_pages, "pdfplumber"


def _iter_docx_blocks(parent):
    """Duyệt theo thứ tự XML: trả về Paragraph và Table xen kẽ.

    Mặc định ``doc.paragraphs`` BỎ QUA mọi bảng — gây mất nội dung HĐLĐ có
    bên/khoản trình bày dạng bảng (Bên B, lương, phụ cấp…). Hàm này lấy đúng
    thứ tự w:p / w:tbl trong body (hoặc trong cell khi đệ quy).
    """
    from docx.document import Document as _Doc
    from docx.oxml.ns import qn
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Doc):
        parent_elem = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elem = parent._tc
    else:
        parent_elem = parent

    for child in parent_elem.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _cell_text(cell) -> str:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parts: list[str] = []
    for blk in _iter_docx_blocks(cell):
        if isinstance(blk, Paragraph):
            t = blk.text.strip()
            if t:
                parts.append(t)
        elif isinstance(blk, Table):
            inner = _table_to_text(blk)
            if inner:
                parts.append(inner)
    return "\n".join(parts).strip()


def _table_to_text(tbl) -> str:
    """Bảng → text. Mỗi hàng `"cell1 | cell2 | ..."`; bỏ ô trùng do merged cell."""
    lines: list[str] = []
    for row in tbl.rows:
        seen: set[str] = set()
        cells_text: list[str] = []
        for cell in row.cells:
            t = _cell_text(cell)
            if not t:
                continue
            if t in seen:
                continue
            seen.add(t)
            cells_text.append(t)
        if cells_text:
            lines.append(" | ".join(cells_text))
    return "\n".join(lines)


def _load_docx(path: Path) -> tuple[str, list[dict[str, str | int]], int | None]:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(path)
    blocks: list[str] = []
    for blk in _iter_docx_blocks(doc):
        if isinstance(blk, Paragraph):
            t = blk.text.strip()
            if t:
                blocks.append(t)
        elif isinstance(blk, Table):
            t = _table_to_text(blk)
            if t:
                blocks.append(t)

    text = "\n".join(blocks)
    page_records = [{"page_num": 1, "text": text}]
    return text, page_records, None


def _load_txt(path: Path) -> tuple[str, list[dict[str, str | int]], int | None]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return raw, [{"page_num": 1, "text": raw}], None


def load(
    file_path: str,
    *,
    pdf_force_ocr: bool = False,
    pdf_detect_scan: bool = True,
) -> ContractDocument:
    """Đọc file và trả ``ContractDocument``.

    PDF scan: nếu ``pdf_force_ocr`` hoặc heuristic ít chữ từ text layer → PaddleOCR (cần ``uv sync --extra ocr``).
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    extraction = "txt"
    if suffix == ".pdf":
        raw, pages, total_pages, extraction = _load_pdf(
            path,
            force_ocr=pdf_force_ocr,
            detect_scan=pdf_detect_scan,
        )
    elif suffix == ".docx":
        raw, pages, total_pages = _load_docx(path)
        extraction = "docx"
    elif suffix in {".txt", ".text"}:
        raw, pages, total_pages = _load_txt(path)
        extraction = "txt"
    else:
        msg = f"Định dạng chưa hỗ trợ: {suffix}. Dùng .pdf, .docx hoặc .txt."
        raise ValueError(msg)

    raw = re.sub(r"[ \t\r\f\v]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw).strip()

    score = score_labor_keywords(raw)
    meta = ContractMetadata(
        filename=path.name,
        total_pages=total_pages,
        detected_language="vi",
        contract_type="labor" if score >= 0.6 else "unknown",
        labor_keyword_score=score,
        extraction_method=extraction,
    )
    return ContractDocument(raw_text=raw, pages=pages, metadata=meta)


def contract_document_from_text(text: str, *, filename: str = "contract.txt") -> ContractDocument:
    """Tạo ``ContractDocument`` từ chuỗi (upload Streamlit)."""
    raw = text.strip()
    score = score_labor_keywords(raw)
    meta = ContractMetadata(
        filename=filename,
        total_pages=None,
        contract_type="labor" if score >= 0.6 else "unknown",
        labor_keyword_score=score,
        extraction_method="plain_upload",
    )
    return ContractDocument(
        raw_text=raw,
        pages=[{"page_num": 1, "text": raw}],
        metadata=meta,
    )
